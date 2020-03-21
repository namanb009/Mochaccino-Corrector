from tkinter import *
from gingerit.gingerit import GingerIt
import docx
import threading
import language_check
import os
from docx.shared import RGBColor
from tkinter import filedialog
from PIL import ImageTk,Image

class Engine():
    add = []  #stores selected file addresses
    filename = []  #stores selected file names
    N=[]
    doc = None  #doc object to work on current word file
    line = []   #stores each line of current file
    n = 0  #number of lines in doc file
    codingfile = False
    docf=None   #Final doc object
   
    
    blocked = ["rape", "terror", "nigga", "nigger", "murder", "suicide", "homicide", "fraud", "drugs", "politics", "earthquake", "tornado",
    "flood", "psycho", "loot", "black man", "black men", "Alcohol", "thief", "war"]
    
    heading = ["Question", "Constraints", "Notes", "Input Format", "Output Format", "Sample Input", "Sample Output", "Explanation"]
    start=[]
    
    def __init__(self):
        abusedoc=docx.Document("abuse.docx")

        for i in range(len(abusedoc.paragraphs)):
            if len(abusedoc.paragraphs[i].text) > 0:
                self.blocked.append(abusedoc.paragraphs[i].text)
    
    
    
    def inputfile(self, root):
        del self.add[::]
        del self.filename[::]
        del self.line[::]

        s = filedialog.askopenfilenames()
        
        for i in s:
            i = i[::-1]
            temp = 0
            
            for j in range(len(i)):
                if i[j] == "/":
                    temp = j
                    break
            
            temp = len(i) - temp
            i = i[::-1]
            
            self.add.append(i[:temp])
            self.filename.append(i[temp:])
        
        self.N=len(self.add)

        l21.config(text=str(self.N))

        temp=""
        for i in self.filename:
            temp=temp+ i + "\n"
        temp=temp[:-1]

        l31.config(text=temp)
   
    def enginecontrol(self):
        for i in range(self.N):
            self.loaddoc(i)

            t1=threading.Thread(target=self.callginger,args=())
            t2 = threading.Thread(target=self.eeoc_checker, args=())
            t3=threading.Thread(target=self.fix_headings,args=())
            
            t1.start()
            t1.join()

            t2.start()
            t3.start()


            t2.join()
            t3.join()

            self.docf.save(self.filename[i][:-5]+"_checked.docx")
    
  
    def loaddoc(self,i):
        del self.line[::]

        os.chdir(self.add[i])

        self.doc = docx.Document(self.filename[i])
        self.doc.save("temp.docx")

        self.doc = docx.Document(self.filename[i])
        self.n = len(self.doc.paragraphs)
            
        for j in range(self.n):
            self.line.append(self.doc.paragraphs[j].text)
            
        self.docf = docx.Document()

    def fix_headings(self):
        found = 0
        for title in self.heading:
            for i in range(self.n):
                if title in self.line[i]:
                    self.start.append(i)
                    self.doc.paragraphs[i].text = str(title)
                    self.doc.paragraphs[i].runs[0].bold = True
                    self.doc.paragraphs[i].runs[0].italic = True
                    self.doc.paragraphs[i].runs[0].underline = True
                    break

            for i in range(self.n):
                if "Sample Case" in self.line[i]:
                    self.doc.paragraphs[i].text = ""    


    def callginger(self):
        for i in range(self.n):
            Ginger = GingerIt()

            if len(self.line[i])>0:

                parastyle = self.doc.paragraphs[i].style
                paraalign = self.doc.paragraphs[i].alignment
                parabold=self.doc.paragraphs[i].runs[0].bold
                paraitalic = self.doc.paragraphs[i].runs[0].italic
                paraunderline = self.doc.paragraphs[i].runs[0].underline
                parafont=self.doc.paragraphs[i].runs[0].text
        
                gin_dic = Ginger.parse(self.line[i])
                gin_dic["corrections"]=gin_dic["corrections"][::-1]     #Reversing dictionary
            
                if len(gin_dic["corrections"]) > 0:
                    if gin_dic["corrections"][0]["start"] == 0:
                        paraObj = self.docf.add_paragraph("")
                    else:
                        paraObj = self.docf.add_paragraph(self.doc.paragraphs[i].text[0: gin_dic["corrections"][0]["start"]])
                
                    for j in range(len(gin_dic["corrections"])):
                        index = gin_dic["corrections"][j]["start"]
                        word = gin_dic["corrections"][j]["correct"]
                    
                        run = paraObj.add_run(gin_dic["result"][index: index + len(word)])
        
                        run.font.color.rgb = RGBColor(0, 100, 0)
                    
                        if j!=len(gin_dic["corrections"])-1:
                            paraObj.add_run(gin_dic["result"][index + len(word) : gin_dic["corrections"][j + 1]["start"]])

                        else:
                            run = paraObj.add_run(gin_dic["result"][index + len(word) ::])
                        


                else:
                    paraObj = self.docf.add_paragraph(self.doc.paragraphs[i].text)

                    
                #Formatting done here
                paraObj.style = parastyle
                paraObj.alignment=paraalign
                    
                for run in paraObj.runs:
                    run.font.bold = parabold
                    run.font.italic = paraitalic
                    run.font.underline = paraunderline
                
            
            else:
                self.docf.add_paragraph("")


    def eeoc_checker(self):
        for i in range(self.n):
            input_term=[]
            input_term = self.line[i].split(" ")

            removeind = []
            removeword=[]

            for j in range(len(input_term)):
                for k in self.blocked:
                    if k in input_term[j]:
                        index=0
                        for word in range(j):
                            index=index+len(input_term[word])+1

                        removeind.append(index)
                        removeword.append(k)

            if len(removeind) > 0:
                self.docf.paragraphs[i].text = ""
                
                if removeind[0] == 0:
                    paraobj = self.docf.paragraphs[i]
                else:
                    paraobj = self.docf.paragraphs[i]
                    paraobj.add_run(self.line[i][0: removeind[0]])
                    
                for ind in range(len(removeind)):
                    paraobj.add_run(self.line[i][removeind[ind]: removeind[ind] + len(removeword[ind])]).font.color.rgb = RGBColor(100, 0, 0)
                    
                    if ind != len(removeind) - 1:
                        paraobj.add_run(self.line[i][removeind[ind] + len(removeword[ind]) : removeind[ind + 1]])
                    else:
                        paraobj.add_run(self.line[i][removeind[ind] + len(removeword[ind]) ::])

#Tkinter window GUI
e=Engine()
root = Tk()
root.title("Mochaccino Grammar Checker")
root.geometry("800x400")

f1 = Frame(root, width=200, height=400)
f1.config(background= "#b7e9fd")
f1.pack(side=LEFT)

f2 = Frame(root, width=400, height=400)
f2.pack()

l0=Label(f2,text="")
l0.grid(row=0,columnspan=2)


b1 = Button(f2, text="Browse files", command=lambda: e.inputfile(root), width=40, height=1)
b1.grid(row=1,column=0,padx=2,pady=2,sticky=W,columnspan=2)

l20=Label(f2,text="Number of files selected: ")
l20.grid(row=2,column=0,padx=2,pady=2,sticky=W)

l21=Label(f2,text="")
l21.grid(row=2,column=1,padx=2,pady=2,sticky=W)

l30=Label(f2,text="File Selected:")
l30.grid(row=3,column=0,padx=2,pady=2,sticky=NW)

l31=Label(f2,text="")
l31.grid(row=3, column=1, padx=2, pady=2, sticky=W)


b41=Button(f2,text="Go!",command=lambda:e.enginecontrol(),width=40,height=1)
b41.grid(row=4, column=0, padx=2, pady=2, sticky=W, columnspan=2)
b41.config(relief="solid")

l50 = Label(f2, text="")
l50.grid(row=5, column=0, padx=2, pady=2, sticky=W)


img = ImageTk.PhotoImage(Image.open("./IMocha.png"))
imgLabel=Label(f2,image=img).grid(row=6,column=0,padx=5,pady=5,columnspan=2)


root.mainloop()